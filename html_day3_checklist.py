from docx import Document

# Tạo tài liệu Word
doc = Document()
doc.add_heading('📋 HTML Day 3 – Daily Learning Report', level=1)

# Danh sách checklist
checklist_day3 = [
    ("1. Hiểu CSS là gì và cách viết CSS", "", ""),
    ("2. Biết 3 cách viết CSS (inline, internal, external)", "", ""),
    ("3. Thực hành đổi màu, font, căn lề bằng CSS", "", ""),
    ("4. Biết cách dùng `class` và `id`", "", ""),
    ("5. Làm lại CV có CSS đẹp hơn", "", ""),
    ("6. Ghi chú lỗi gặp phải khi dùng CSS", "", ""),
    ("7. Đánh giá mức độ hiểu bài (1–5 ⭐)", "", ""),
    ("8. Lên kế hoạch Ngày 4", "", "Ví dụ: học Flexbox, Grid...")
]

# Tạo bảng
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nội dung'
hdr_cells[1].text = 'Hoàn thành (✔️/❌)'
hdr_cells[2].text = 'Ghi chú'

for item in checklist_day3:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]

# Lưu file Word
doc.save("HTML_Day3_Checklist.docx")
print("✅ Đã tạo file: HTML_Day3_Checklist.docx")
